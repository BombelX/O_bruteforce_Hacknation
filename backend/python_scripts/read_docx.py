import os
from docx import Document
def docx_to_table(docx_path):
    print(f"Otwieranie pliku DOCX: ",docx_path)

    if not os.path.exists(docx_path):
        print(f"Błąd: Plik {docx_path} nie istnieje.")
        return [],0,0
    try:
        doc=Document(docx_path)
        if len(doc.tables)==0:
            print("Nie znaleziono żadnej tabeli w pliku DOCX.")
            return [],0,0

        table = doc.tables[0]

        parsed_data=[]
        rows=table.rows
        print(f"Znaleziono {len(rows)} wierszy w tabeli.")

        for row in rows:
           cols = [cell.text.strip() for cell in row.cells]
           if any(cols):
                parsed_data.append(cols)
        if not parsed_data:
            return [],0,0

        columns_no=len(parsed_data[0])
        rows_no=len(parsed_data)
        return [parsed_data],columns_no,rows_no

    except Exception as e:
        print(f"Błąd podczas przetwarzania pliku DOCX: {e}")
        return [],0,0
